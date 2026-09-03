"""Builds the 8 bundled default .docx templates (4 document types x
English/Arabic) that ship under backend/app/assets/doc_templates/ --
see doc_template_service.py, which falls back to these whenever an
admin hasn't uploaded a custom template for a given (doc_type,
language) pair.

Run from the backend/ directory (needs the venv's python-docx):

    python scripts/generate_default_doc_templates.py

Safe to re-run any time the defaults need tweaking -- it always
overwrites all 8 files with a fresh build from the definitions below,
there's no partial/incremental state to worry about.

Templates use docxtpl's Jinja2-in-Word syntax: {{ field }} for a plain
substitution, and {%tr for line in lines %} ... {%tr endfor %} to
repeat a table row once per line item (docxtpl's row-repeat tag --
plain Jinja {% %} can't repeat a *row*, only text within one). Every
tag here is written as a single python-docx run (`paragraph.add_run` /
`cell.text =`), never split across runs -- exactly the failure mode
that makes hand-typing a tag in Word risky (AutoCorrect/spell-check can
silently split a run mid-tag), so a template rebuilt from this script
is always safe to open in docxtpl even before anyone edits it further.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

OUT_DIR = Path(__file__).resolve().parent.parent / "app" / "assets" / "doc_templates"

HEADER_FILL = "2D3748"


def _set_rtl(paragraph) -> None:
    """python-docx has no high-level right-to-left API -- this sets the
    paragraph's reading order directly via its underlying XML (<w:bidi/>
    on <w:pPr>), which is what Word actually checks to lay Arabic text
    out right-to-left instead of just rendering Arabic glyphs
    left-to-right."""
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.makeelement(qn("w:bidi"), {})
    p_pr.append(bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def _set_cell_rtl(cell) -> None:
    for paragraph in cell.paragraphs:
        _set_rtl(paragraph)


def _add_run(paragraph, text: str, *, bold: bool = False, size: int | None = None, rtl: bool = False, color=None):
    run = paragraph.add_run(text)
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if rtl:
        # Complex-script font: Word picks a fallback for Arabic glyphs
        # automatically, but setting it explicitly (rather than leaving
        # only the Latin-script rFonts populated) keeps rendering
        # consistent across machines that don't share the same default
        # fallback chain.
        rpr = run._r.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = rpr.makeelement(qn("w:rFonts"), {})
            rpr.append(rfonts)
        rfonts.set(qn("w:cs"), "Arial")
        rtl_el = rpr.makeelement(qn("w:rtl"), {})
        rpr.append(rtl_el)
    return run


def _shade_cell(cell, color_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.makeelement(qn("w:shd"), {qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): color_hex})
    tc_pr.append(shd)


LABELS = {
    "en": {
        "quotation": "QUOTATION",
        "order": "SALES ORDER",
        "delivery_note": "DELIVERY NOTE",
        "feasibility": "FEASIBILITY REPORT",
        "customer": "Customer",
        "date": "Date",
        "quotation_date": "Quotation Date",
        "valid_until": "Valid Until",
        "order_date": "Order Date",
        "requested_delivery": "Requested Delivery",
        "delivery_date": "Delivery Date",
        "against_order": "Against Order",
        "required_by": "Required By",
        "checked_at": "Checked At",
        "deal": "Deal",
        "status": "Status",
        "no": "#",
        "product": "Product",
        "qty": "Qty",
        "qty_delivered": "Qty Delivered",
        "unit": "Unit",
        "unit_price": "Unit Price",
        "line_total": "Line Total",
        "supply_plan": "Supply Plan",
        "materials": "Materials",
        "subtotal": "Subtotal",
        "discount": "Discount",
        "total": "Total",
        "notes": "Notes",
        "signature": "Authorized Signatory",
        "generated": "Generated",
    },
    "ar": {
        "quotation": "عرض سعر",
        "order": "أمر بيع",
        "delivery_note": "إشعار تسليم",
        "feasibility": "تقرير الجدوى",
        "customer": "العميل",
        "date": "التاريخ",
        "quotation_date": "تاريخ العرض",
        "valid_until": "صالح حتى",
        "order_date": "تاريخ الطلب",
        "requested_delivery": "تاريخ التسليم المطلوب",
        "delivery_date": "تاريخ التسليم",
        "against_order": "مقابل الطلب رقم",
        "required_by": "مطلوب بحلول",
        "checked_at": "تاريخ الفحص",
        "deal": "الصفقة",
        "status": "الحالة",
        "no": "#",
        "product": "المنتج",
        "qty": "الكمية",
        "qty_delivered": "الكمية المسلمة",
        "unit": "الوحدة",
        "unit_price": "سعر الوحدة",
        "line_total": "إجمالي السطر",
        "supply_plan": "خطة التوريد",
        "materials": "المواد",
        "subtotal": "المجموع الفرعي",
        "discount": "الخصم",
        "total": "الإجمالي",
        "notes": "ملاحظات",
        "signature": "التوقيع المعتمد",
        "generated": "تاريخ الإصدار",
    },
}


def _build(*, language: str, title_key: str, doc_number_field: str, meta_fields, line_columns, has_total: bool):
    rtl = language == "ar"
    L = LABELS[language]
    doc = Document()

    section = doc.sections[0]
    section.left_margin = section.right_margin = Mm(18)
    section.top_margin = section.bottom_margin = Mm(16)

    # --- Letterhead ------------------------------------------------------
    header = doc.add_table(rows=1, cols=2)
    header.autofit = True
    left, right = header.rows[0].cells
    p = left.paragraphs[0]
    if rtl:
        _set_rtl(p)
    _add_run(p, "{{ company_name }}", bold=True, size=16, rtl=rtl)

    p2 = right.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT if not rtl else WD_ALIGN_PARAGRAPH.RIGHT
    _add_run(p2, L[title_key], bold=True, size=14, rtl=rtl)
    p3 = right.add_paragraph()
    if rtl:
        _set_rtl(p3)
    _add_run(p3, "{{ " + doc_number_field + " }}", rtl=rtl)

    doc.add_paragraph()

    # --- Meta / party block -----------------------------------------------
    meta_table = doc.add_table(rows=len(meta_fields), cols=2)
    meta_table.autofit = True
    for row, (label_key, placeholder) in zip(meta_table.rows, meta_fields):
        label_cell, value_cell = row.cells
        lp = label_cell.paragraphs[0]
        vp = value_cell.paragraphs[0]
        if rtl:
            _set_rtl(lp)
            _set_rtl(vp)
        _add_run(lp, L[label_key] + ":", bold=True, rtl=rtl)
        _add_run(vp, placeholder, rtl=rtl)

    doc.add_paragraph()

    # --- Line items table ---------------------------------------------------
    # docxtpl's row-repeat tag needs each of {%tr for %} / {%tr endfor %}
    # to be the *sole* content of its own row -- docxtpl's XML
    # preprocessing collapses such a row down to bare Jinja text (see
    # backend/scripts/generate_default_doc_templates.py's module
    # docstring), so the for/endfor markers vanish and only the data row
    # between them repeats. Putting the tags inside the data row's own
    # cells (alongside real content) instead silently eats the whole row
    # -- confirmed against docxtpl 0.20.2's actual preprocessing regex,
    # not just the tag names.
    n_cols = len(line_columns)
    table = doc.add_table(rows=4, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    header_row, for_row, data_row, endfor_row = table.rows

    for idx, (label_key, _tpl) in enumerate(line_columns):
        cell = header_row.cells[idx]
        _shade_cell(cell, HEADER_FILL)
        p = cell.paragraphs[0]
        if rtl:
            _set_rtl(p)
        _add_run(p, L[label_key], bold=True, color="FFFFFF", rtl=rtl)

    for_row.cells[0].text = "{%tr for line in lines %}"
    endfor_row.cells[0].text = "{%tr endfor %}"

    for idx, (_label_key, tpl) in enumerate(line_columns):
        cell = data_row.cells[idx]
        p = cell.paragraphs[0]
        if rtl:
            _set_rtl(p)
        _add_run(p, tpl, rtl=rtl)

    doc.add_paragraph()

    # --- Totals (quotation/order only) -------------------------------------
    if has_total:
        totals = doc.add_table(rows=3, cols=2)
        totals.alignment = WD_TABLE_ALIGNMENT.RIGHT
        rows_def = [("subtotal", "{{ subtotal_amount }}"), ("discount", "{{ discount_amount }}"), ("total", "{{ total_amount }}")]
        for row, (label_key, placeholder) in zip(totals.rows, rows_def):
            label_cell, value_cell = row.cells
            lp = label_cell.paragraphs[0]
            vp = value_cell.paragraphs[0]
            lp.alignment = WD_ALIGN_PARAGRAPH.RIGHT if not rtl else WD_ALIGN_PARAGRAPH.LEFT
            vp.alignment = WD_ALIGN_PARAGRAPH.RIGHT if not rtl else WD_ALIGN_PARAGRAPH.LEFT
            if rtl:
                _set_rtl(lp)
                _set_rtl(vp)
            bold = label_key == "total"
            _add_run(lp, L[label_key] + ":", bold=bold, rtl=rtl)
            _add_run(vp, placeholder, bold=bold, rtl=rtl)
        doc.add_paragraph()

    # --- Notes ---------------------------------------------------------------
    notes_p = doc.add_paragraph()
    if rtl:
        _set_rtl(notes_p)
    _add_run(notes_p, L["notes"] + ": ", bold=True, rtl=rtl)
    _add_run(notes_p, "{{ notes }}", rtl=rtl)

    doc.add_paragraph()
    doc.add_paragraph()

    # --- Signature block -------------------------------------------------
    sig_p = doc.add_paragraph()
    if rtl:
        _set_rtl(sig_p)
    _add_run(sig_p, "{{ signature_image }}", rtl=rtl)

    line_p = doc.add_paragraph("____________________")
    if rtl:
        _set_rtl(line_p)

    name_p = doc.add_paragraph()
    if rtl:
        _set_rtl(name_p)
    _add_run(name_p, "{{ signer_name }}", rtl=rtl)

    role_p = doc.add_paragraph()
    if rtl:
        _set_rtl(role_p)
    _add_run(role_p, L["signature"], size=9, rtl=rtl)

    footer_p = doc.add_paragraph()
    if rtl:
        _set_rtl(footer_p)
    _add_run(footer_p, L["generated"] + ": {{ generated_date }}", size=8, rtl=rtl)

    return doc


DEFS = {
    "quotation": dict(
        title_key="quotation",
        doc_number_field="quotation_number",
        meta_fields=[
            ("customer", "{{ customer_name }}"),
            ("quotation_date", "{{ quotation_date }}"),
            ("valid_until", "{{ valid_until }}"),
            ("status", "{{ status_label }}"),
        ],
        line_columns=[
            ("no", "{{ line.index }}"),
            ("product", "{{ line.product_code }} - {{ line.product_name }}"),
            ("qty", "{{ line.quantity }}"),
            ("unit", "{{ line.unit }}"),
            ("unit_price", "{{ line.unit_price }}"),
            ("line_total", "{{ line.line_total }}"),
        ],
        has_total=True,
    ),
    "order": dict(
        title_key="order",
        doc_number_field="order_number",
        meta_fields=[
            ("customer", "{{ customer_name }}"),
            ("order_date", "{{ order_date }}"),
            ("requested_delivery", "{{ requested_delivery_date }}"),
            ("status", "{{ status_label }}"),
        ],
        line_columns=[
            ("no", "{{ line.index }}"),
            ("product", "{{ line.product_code }} - {{ line.product_name }}"),
            ("qty", "{{ line.quantity }}"),
            ("unit", "{{ line.unit }}"),
            ("unit_price", "{{ line.unit_price }}"),
            ("line_total", "{{ line.line_total }}"),
        ],
        has_total=True,
    ),
    "delivery_note": dict(
        title_key="delivery_note",
        doc_number_field="delivery_note_number",
        meta_fields=[
            ("customer", "{{ customer_name }}"),
            ("delivery_date", "{{ delivery_date }}"),
            ("against_order", "{{ order_number }}"),
            ("status", "{{ status_label }}"),
        ],
        line_columns=[
            ("no", "{{ line.index }}"),
            ("product", "{{ line.product_code }} - {{ line.product_name }}"),
            ("qty_delivered", "{{ line.quantity_delivered }}"),
            ("unit", "{{ line.unit }}"),
        ],
        has_total=False,
    ),
    "feasibility": dict(
        title_key="feasibility",
        doc_number_field="feasibility_number",
        meta_fields=[
            ("customer", "{{ customer_name }}"),
            ("deal", "{{ deal_number }}"),
            ("required_by", "{{ required_by_date }}"),
            ("checked_at", "{{ checked_at }}"),
            ("status", "{{ status_label }}"),
        ],
        line_columns=[
            ("no", "{{ line.index }}"),
            ("product", "{{ line.product_code }} - {{ line.product_name }}"),
            ("qty", "{{ line.quantity }}"),
            ("supply_plan", "{{ line.supply_note }}"),
            ("materials", "{{ line.feasible_label }}"),
        ],
        has_total=False,
    ),
}


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    for doc_type, definition in DEFS.items():
        for language in ("en", "ar"):
            doc = _build(language=language, **definition)
            out_path = OUT_DIR / f"{doc_type}_{language}.docx"
            doc.save(out_path)
            print(f"wrote {out_path}")


if __name__ == "__main__":
    main()
